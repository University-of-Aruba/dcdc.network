from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_travel_request(filename, destination, institution, dates_travel, dates_activity, justification, multi_purpose, cost_estimate):
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    title = doc.add_heading('DCDC Travel and Activity Request Form', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('Submit to the DCDC Network Administrator: rendell.dekort@ua.aw')
    doc.add_paragraph()

    # Section 1
    doc.add_heading('Section 1: Applicant details', level=2)
    table1 = doc.add_table(rows=4, cols=2)
    table1.style = 'Table Grid'
    fields1 = [
        ('Name:', 'Rendell de Kort'),
        ('Role in DCDC Network', 'Network Administrator'),
        ('Island / institution', 'Aruba / University of Aruba'),
        ('Date of submission', '14-4-2026')
    ]
    for i, (field, response) in enumerate(fields1):
        table1.rows[i].cells[0].text = field
        table1.rows[i].cells[1].text = response

    doc.add_paragraph()

    # Section 2
    doc.add_heading('Section 2: Activity details', level=2)
    table2 = doc.add_table(rows=7, cols=2)
    table2.style = 'Table Grid'
    fields2 = [
        ('Name of event / activity', 'DCDC "Introduction to R for SPSS Users" - ' + destination + ' delivery'),
        ('Type', 'Training delivery + stakeholder engagement'),
        ('Location', destination + ', ' + institution),
        ('Dates of travel', dates_travel),
        ('Dates of event / activity', dates_activity),
        ('Organizer / host', 'DCDC Network'),
        ('Link to event (if applicable)', 'https://dcdc.network')
    ]
    for i, (field, response) in enumerate(fields2):
        table2.rows[i].cells[0].text = field
        table2.rows[i].cells[1].text = response

    doc.add_paragraph()

    # Section 3
    doc.add_heading('Section 3: Budget and cost estimate', level=2)
    doc.add_paragraph('Budget lines: Training (E150,000 pool) / Additional travel (E65,000 pool)')

    table3 = doc.add_table(rows=len(cost_estimate)+1, cols=3)
    table3.style = 'Table Grid'
    table3.rows[0].cells[0].text = 'Item'
    table3.rows[0].cells[1].text = 'Estimated cost (EUR)'
    table3.rows[0].cells[2].text = 'Budget line'
    for i, (item, cost, budget) in enumerate(cost_estimate):
        table3.rows[i+1].cells[0].text = item
        table3.rows[i+1].cells[1].text = cost
        table3.rows[i+1].cells[2].text = budget

    doc.add_paragraph()

    # Section 4
    doc.add_heading('Section 4: Justification', level=2)

    doc.add_heading('4a. Connection to DCDC Network objectives', level=3)
    doc.add_paragraph('This activity contributes to:')
    for obj in ['Building digital competence and data skills across the region',
                'Peer-to-peer exchange and expertise sharing',
                'Training program development or delivery',
                'Stakeholder engagement and network growth']:
        doc.add_paragraph(obj, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Description: ' + justification)

    doc.add_heading('4b. FAIR principles and open science alignment', level=3)

    p1 = doc.add_paragraph()
    run1 = p1.add_run('Will insights, materials, or outputs from this activity be shared with the broader network?\n')
    run1.bold = True
    p1.add_run('Yes. All course materials are developed as open educational resources hosted on GitHub (University-of-Aruba organization) and will be submitted to the Carpentries Incubator. Materials are licensed for reuse. The course uses open-source software (R, RStudio) to replace proprietary tools (SPSS). Post-delivery, updated materials reflecting island-specific feedback will be pushed to the shared repository.')

    doc.add_paragraph()
    p2 = doc.add_paragraph()
    run2 = p2.add_run('Does this activity involve or promote open data, open access publishing, or FAIR data practices?\n')
    run2.bold = True
    p2.add_run('Yes, directly. The course teaches participants to work with FAIR-compliant data workflows in R, including reproducible analysis scripts, transparent data import and cleaning, and version-controlled project structures. This is a core component of the DCDC Network open science mission.')

    doc.add_paragraph()
    p3 = doc.add_paragraph()
    run3 = p3.add_run('Are there any outputs from this activity that could be made openly available with a persistent identifier?\n')
    run3.bold = True
    p3.add_run('Yes. The course materials will be archived on Zenodo with a DOI after the pilot cycle is complete. Participant feedback summaries (anonymized) will also be archived as part of the train-the-trainer documentation.')

    doc.add_paragraph()
    doc.add_heading('4c. Multi-purpose value', level=3)
    p4 = doc.add_paragraph()
    run4 = p4.add_run('Does this trip combine multiple network objectives?\n')
    run4.bold = True
    p4.add_run(multi_purpose)

    doc.add_paragraph()
    doc.add_heading('4d. Timing and urgency', level=3)
    doc.add_paragraph('No, standard timeline applies. Exact dates to be confirmed with the host institution. Submission at this stage is to ensure the trip is budgeted and approved in advance per the SOP.')

    doc.add_paragraph()

    # Section 5
    doc.add_heading('Section 5: Post-activity report (to be completed after the activity)', level=2)
    doc.add_paragraph('Submit to the administrator within four weeks of returning.')

    table5 = doc.add_table(rows=5, cols=2)
    table5.style = 'Table Grid'
    for i, field in enumerate(['Summary of activity attended',
                               'Key insights relevant to the DCDC Network',
                               'Planned follow-up actions',
                               'Materials or outputs to be shared with the network',
                               'Actual costs (attach receipts)']):
        table5.rows[i].cells[0].text = field
        table5.rows[i].cells[1].text = ''

    doc.add_paragraph()
    doc.add_paragraph('End of form')

    doc.save(filename)
    print(f"Saved: {filename}")


# Curacao
create_travel_request(
    filename=r"C:\Users\Rendell CE\Documents\GitHub\knowledge-hub\01-projects\dcdc-travel\DCDC Travel Request - Rendell de Kort - Curacao R Course.docx",
    destination="Curacao",
    institution="University of Curacao",
    dates_travel="TBC (August-September 2026, estimated 4-5 days)",
    dates_activity="TBC (2 course days within travel window)",
    justification="Delivery of the DCDC Network flagship training course, Introduction to R for SPSS Users, at the University of Curacao. This is the second island delivery following the Aruba pilot (April 2026). The course replaces proprietary SPSS with open-source R across Dutch Caribbean research institutions, directly advancing the network digital competence and open science objectives. The course was developed collaboratively with input from all three island coordinators and is being submitted to the Carpentries Incubator for broader adoption.",
    multi_purpose="Yes. In addition to the two-day course delivery, this trip will include: (1) coordination meeting with the Curacao local coordinator (Marjorie Alfonso) on network operations, event planning, and the November kickoff conference; (2) stakeholder engagement with UoC faculty and researchers to assess demand for follow-up training (e.g., intermediate R, data visualization); (3) train-the-trainer handoff session to build local capacity for future course delivery without administrator travel.",
    cost_estimate=[
        ('Return flights (AUA-CUR)', '200-350', 'Additional travel'),
        ('Accommodation (3-4 nights)', '300-500', 'Additional travel'),
        ('Local transport', '50-100', 'Additional travel'),
        ('Per diem', '160-240', 'Additional travel'),
        ('Course materials / printing', '50-100', 'Training'),
        ('Total estimate', '760-1,290', ''),
    ]
)

# Sint Maarten
create_travel_request(
    filename=r"C:\Users\Rendell CE\Documents\GitHub\knowledge-hub\01-projects\dcdc-travel\DCDC Travel Request - Rendell de Kort - Sint Maarten R Course.docx",
    destination="Sint Maarten",
    institution="University of St. Martin",
    dates_travel="TBC (September-October 2026, estimated 4-5 days)",
    dates_activity="TBC (2 course days within travel window)",
    justification="Delivery of the DCDC Network flagship training course, Introduction to R for SPSS Users, at the University of St. Martin. This is the third and final island delivery in the pilot cycle, following Aruba (April 2026) and Curacao (August-September 2026). By this point, the course materials will have been refined based on two prior deliveries, and the train-the-trainer methodology will be tested. The Sint Maarten delivery completes the network first full training cycle across all three CAS universities.",
    multi_purpose="Yes. In addition to the two-day course delivery, this trip will include: (1) coordination meeting with the Sint Maarten local coordinator on network operations and the November kickoff conference (Nov 10-12); (2) stakeholder engagement with USM faculty and IPA (Institute for Professional Advancement) on follow-up training needs; (3) train-the-trainer session to ensure local delivery capacity post-pilot; (4) if timing aligns, preliminary logistics coordination for the DCDC Kickoff Conference (Nov 10-12, location TBC).",
    cost_estimate=[
        ('Return flights (AUA-SXM)', '300-500', 'Additional travel'),
        ('Accommodation (3-4 nights)', '400-600', 'Additional travel'),
        ('Local transport', '50-100', 'Additional travel'),
        ('Per diem', '160-240', 'Additional travel'),
        ('Course materials / printing', '50-100', 'Training'),
        ('Total estimate', '960-1,540', ''),
    ]
)

print("\nBoth forms created.")
